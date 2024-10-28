# Framework supporting MLOps Apps
# %%

import os
import time
import datetime
import streamlit as st
from PIL import Image
from streamlit_extras.app_logo import add_logo
from langchain_community.llms import OpenAI
from llama_index.llms.azure_openai import AzureOpenAI
from llama_index.embeddings.azure_openai import AzureOpenAIEmbedding
from llama_index.core import VectorStoreIndex, SimpleDirectoryReader
from llama_index.core.service_context import ServiceContext
from langchain.schema import SystemMessage, HumanMessage, AIMessage
from dotenv import load_dotenv
import streamlit.components.v1 as components
import docx
import base64
import pandas as pd
import numpy as np

# Load environment variables
load_dotenv(override=True)
OPENAI_KEY = os.getenv('OPENAI_API_KEY')
ENDPOINT = os.getenv('OPENAI_API_BASE')
OPENAI_VERSION = os.getenv('OPENAI_API_VERSION')

# Sidebar for navigation
with st.sidebar:
    selected = st.selectbox('OptiLabor Tool', 
                             ['Virtual Assistant',
                              'LaborSync: Automation (Augmented Reality + Digital Twins + On-Site Robots)',
                              'Upskilling Module: Training & Competency'],
                             index=0)

# Front-End Layout
st.set_page_config(page_title='OptiLabor AI', page_icon='OptiLaborAI.png', layout="wide")
add_logo("OptiLaborAI.png")
st.title("💸 OptiLabor AI")
st.markdown("Your AI assistant for construction projects, providing real-time information, collaboration, and resource predictions.")

# Display features
st.markdown("### Features")
features = [
    "OL Takeoff - Submit your blueprint for an accurate Bill of Quantities.",
    "OL Doc Wizard - Quickly find invoices, contracts, and create reports.",
    "OL Virtual Design - Identify design flaws early with 3D models.",
    "OL Project Management - Gain insights from real-time data using AI and IoT.",
    "OL Decision Making - Visualize KPIs and automate compliance reporting."
]
st.markdown("\n".join(f"- {feature}" for feature in features))
st.markdown("-------")

# Display video and image
st.video(open('Untitled video.mp4', 'rb').read())
st.image('OptiLabor.png', caption='OptiLabor Features')

# Virtual Assistant Module
if selected == 'Virtual Assistant':
    st.title("Module 1: Virtual Assistant")
    st.markdown("An NLP assistant providing real-time information, reducing RFIs.")
    st.header("Document Management Challenges")
    st.image('docs construction.jpeg', caption='Source: InEight Blog')

    def init_page():
        st.header('OL DocWizard')
        st.write("I'm here to help you get information from your database documents.")
        st.sidebar.title('Options')

    def select_llm():
        return AzureOpenAI(model='gpt-35-turbo-16k', deployment_name='GPT35-optilabor', 
                           api_key=OPENAI_KEY, azure_endpoint=ENDPOINT, api_version=OPENAI_VERSION)

    def select_embedding():
        return AzureOpenAIEmbedding(model='text-embedding-ada-002', deployment_name='text-embedding-ada-002', 
                                     api_key=OPENAI_KEY, azure_endpoint=ENDPOINT, api_version=OPENAI_VERSION)

    def init_messages():
        if 'messages' not in st.session_state or st.sidebar.button('Clear Conversation'):
            st.session_state.messages = [SystemMessage(content='You are a helpful AI assistant. Reply in markdown format.')]

    def get_answer(query_engine, messages):
        response = query_engine.query(messages)
        return response.response, response.metadata

    def main():
        init_page()
        file = st.file_uploader('Upload file:', type=['pdf', 'txt', 'docx'])
        
        if file:
            with open(os.path.join('data', file.name), 'wb') as f: 
                f.write(file.getbuffer())        
            st.success('Saved file!')

            documents = SimpleDirectoryReader('./data').load_data()
            file_names = [doc.metadata['file_name'] for doc in documents]
            st.write('Current documents in folder:', ', '.join(file_names))

            llm = select_llm()
            embed = select_embedding()
            service_context = ServiceContext.from_defaults(llm=llm, embed_model=embed)
            query_engine = VectorStoreIndex.from_documents(documents, service_context=service_context).as_query_engine()

            init_messages()

            if user_input := st.chat_input('Input your question!'):
                st.session_state.messages.append(HumanMessage(content=user_input))
                with st.spinner('Bot is typing ...'):
                    answer, meta_data = get_answer(query_engine, user_input)

                greetings = [...]  # Define greeting phrases
                compliments = [...]  # Define compliment phrases
                if user_input.lower() in greetings:
                    answer = 'Hi, how can I assist you?'
                elif user_input.lower() in compliments:
                    answer = 'My pleasure! Feel free to ask more questions.'
                elif any(keyword in answer for keyword in keywords):  # Define keywords
                    st.session_state.messages.append(AIMessage(content=f"**Source**: {list(meta_data.values())[0]['file_name']}  \n**Answer**: {answer}"))
                else:
                    answer = 'This is outside the scope of the provided knowledge base.'

                st.session_state.messages.append(AIMessage(content=answer))

            for message in st.session_state.get('messages', []):
                with st.chat_message('assistant' if isinstance(message, AIMessage) else 'user'):
                    st.markdown(message.content)
        else:
            if not os.listdir('./data'):
                st.write('No file is saved yet.')

    if __name__ == '__main__':
        main()

    # Reporting Visualization
    st.title("A lot of data is wasted")
    st.header("Construction projects generate massive amounts of data, yet 80% of it remains unstructured.")
    st.image('reporting.png', caption='Source: StructionSite Blog')

    # Show Power BI report
    path_to_html = "/Users/juanrivera/Desktop/chatbot1/Streamlit/pages/11_power_BI.html"
    with open(path_to_html, 'r') as f: 
        components.html(f.read(), height=2000)

# Chatbot Report Generation
st.markdown("<h1 style='text-align:justified;font-family:Georgia'>Construction Chatbot - Doc Generator</h1>", unsafe_allow_html=True)
with st.sidebar:
    openai_api_key = st.secrets["auth_token"]
    st.markdown("-------")
    company_name = st.text_input("What is the name of the company?")
    start_up_description = st.text_input("Please describe your statement and objectives of your report")
    sector = st.multiselect('In which sector of construction is the project?', ["Non-Residential", "Residential"])
    st.markdown("-------")
    generate_button = st.button("Generate my Report")

def generate_report(company_name, report_date):
    doc = docx.Document()
    doc.add_heading("Report", 0)
    doc.add_paragraph(f'Created On: {report_date}')
    doc.add_paragraph(f'Created For: {company_name}')
    doc.add_heading(f'Balance of {company_name} for {", ".join(sector)} sector')
    doc.save('Construction Report.docx')
    
    with open('Construction Report.docx', "rb") as file:
        data = file.read()
        encoded = base64.b64encode(data).decode('utf-8')
    st.download_button('Download Here', encoded, "Construction Report.docx")

def generate_response(input_text):
    llm = OpenAI(temperature=0.3, openai_api_key=openai_api_key)
    return llm(input_text)

if "messages" not in st.session_state:
    st.session_state.messages = [{"role": "", "content": "Hey there, I'm OptiLabor Bot, here to help you create your report. Please input your report details on the left."}]

for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

if generate_button:
    if openai_api_key.startswith('sk-'):
        date_today = datetime.date.today()
        funding_summary = generate_response(f"I'm exploring funding options for my startup '{company_name}'. The description is: {start_up_description}. Please provide an overview of different funding sources available to early-stage startups in the {', '.join(sector)} sector.")
        legal_summary = generate_response(f"I need legal guidance for launching '{company_name}' in the {', '.join(sector)} sector. Please provide relevant legal requirements and regulations.")
        
        generate_report(company_name, date_today)
    else:
        st.warning('Please enter your OpenAI API key!', icon='⚠')

if (prompt := st.chat_input("What is up?")): 
    if openai_api_key.startswith('sk-'):
        st.session_state.messages.append({"role": "", "content": prompt})
        
        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            message_placeholder = st.empty()
            full_response = ""
            with st.spinner('Wait for it...'):
                assistant_response = generate_response(prompt)
                for chunk in assistant_response.split():
                    full_response += chunk + " "
                    time.sleep(0.05)
                    message_placeholder.markdown(full_response + "▌")
            message_placeholder.markdown(full_response)
            st.session_state.messages.append({"role": "", "content": full_response})
    else:
        st.warning('Please enter your OpenAI API key!', icon='⚠')



